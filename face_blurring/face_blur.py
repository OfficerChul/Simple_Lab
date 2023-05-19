import os
import cv2
from retinaface import RetinaFace
import numpy as np
import matplotlib.pyplot as plt
from datetime import timedelta
import time
import docx

if os.path.exists(r'/root/Simple_Lab/face_blurring/progress.docx'):
    os.remove(r'/root/Simple_Lab/face_blurring/progress.docx')

try:
    doc = docx.Document(r'/root/Simple_Lab/face_blurring/progress.docx')
except:
    doc = docx.Document()
    doc.save(r'/root/Simple_Lab/face_blurring/progress.docx')
    print('The document created.')


tic = time.time()


# example_folder_path = r'/root/Simple_Lab/face_blurring/example_vids/'
FHD_FWL = r'/kyochul/home/kyochul/FHD_FWL'
out_path = r'/kyochul/home/kyochul/BLURRED_FHD_FWL'

# text file to keep up the progress
# prog_file = open(r'/root/Simple_Lab/face_blurring/progress.txt', "w+")
try:
    vid_num = 0
    for action_camera in os.listdir(FHD_FWL):
        
        action_camera_path = os.path.join(FHD_FWL, action_camera)

        # make directory if not exist
        out_action_camera_path = os.path.join(out_path, action_camera)
        if not os.path.exists(out_action_camera_path):
            os.mkdir(out_action_camera_path)

        for action in os.listdir(action_camera_path):

            action_path = os.path.join(action_camera_path, action)

            # make directory if not exist
            out_action_path = os.path.join(out_action_camera_path, action)
            if not os.path.exists(out_action_path):
                os.mkdir(out_action_path)

            # video start
            for vid in os.listdir(action_path):         
                video_path = os.path.join(action_path, vid)

                out_video_path = os.path.join(out_action_path, vid)

                tic3 = time.time() # tictoc for one video

                cap = cv2.VideoCapture(video_path)
                fourcc = cv2.VideoWriter_fourcc(*'mp4v')

                fps = cap.get(cv2.CAP_PROP_FPS)

                # temp_out_vid_path = os.path.join(r'/root', vid)
                # out = cv2.VideoWriter(temp_out_vid_path, fourcc, fps, (1920, 1080))

                out = cv2.VideoWriter(out_video_path, fourcc, fps, (1920, 1080))

                total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                msg1 = f'{video_path}\ntotal frame: {total}'

                print(msg1)
                # prog_file.write(msg1 + '\n')
                doc.add_paragraph(msg1)
                doc.save(r'/root/Simple_Lab/face_blurring/progress.docx')

                vid_num += 1

                count = 0
                while True:
                    tic2 = time.time()
                    ret, frame = cap.read()
                    if not ret:
                        break
                    count += 1
                    # print(f'count: {count}/{total}')
                    frameRGB = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    faces = RetinaFace.detect_faces(frameRGB)
                    if type(faces) != tuple: # if the face cannot be detected from the video

                        for face in faces.keys():
                            face = faces[face]
                            facial_area = face['facial_area']
                            x = facial_area[0]
                            y = facial_area[1]
                            w = facial_area[2] - facial_area[0]
                            h = facial_area[3] - facial_area[1]
                            roi = frame[y:y + h, x:x + w]

                            # apply gaussian blur to face rectangle
                            roi = cv2.GaussianBlur(roi, (17, 17), 30)

                            # add blurred face on original image to get final image
                            frame[y:y + roi.shape[0], x:x + roi.shape[1]] = roi

                    out.write(frame)
                    if cv2.waitKey(1) & 0xFF == ord('q'):
                        break
                    elapsed2 = time.time() - tic2
                    formatted_elapsed2 = str(timedelta(seconds=elapsed2))
                    print(f'count: {count}/{total}, elapsed time(%hh:%mm:%ss.xxx): {formatted_elapsed2}, progess: {vid_num}/25210')
                    
                cap.release()
                out.release()
                elapsed3 = time.time() - tic3
                formatted_elapsed3 = str(timedelta(seconds=elapsed3))
                msg2 = f'{vid_num}th / 25210 video took {formatted_elapsed3} time elapsed, {round((vid_num / 25210) * 100, 4)}% done\n'
                print(msg2)
                # prog_file.write(msg2 + '\n')
                doc.add_paragraph(msg2)
                doc.add_paragraph('-----------------------------------------')
                doc.save(r'/root/Simple_Lab/face_blurring/progress.docx')

    # prog_file.close()
                

    elapsed = time.time() - tic
    formatted_elapsed = str(timedelta(seconds=elapsed))
    print(f'Total Elapsed Time(%hh:%mm:%ss.xxx): {formatted_elapsed}')

except Exception as error:
    doc.add_paragraph(error)
            
    






